import sqlite3
import tkinter as tk
from tkinter import ttk, messagebox
import datetime
from docx import Document

def create_db():
    conn = sqlite3.connect('forensic_form.db')
    cur = conn.cursor()
    cur.execute('''
        CREATE TABLE IF NOT EXISTS forensic_data (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            surname TEXT,
            name TEXT,
            patronymic TEXT,
            gender TEXT,
            birthdate TEXT,
            passport_number TEXT,
            residence TEXT,
            age TEXT,
            nationality TEXT,
            death_date TEXT,
            body_received_date TEXT,
            body_received_time TEXT,
            body_found_district TEXT,
            body_comment TEXT,
            examiner TEXT,
            examination_start_date TEXT,
            examination_end_date TEXT,
            diagnosis TEXT,
            complications TEXT,
            background_pathology TEXT,
            death_type TEXT,
            death_kind TEXT,
            conclusion_receiver TEXT,
            conclusion_issue_date TEXT,
            death_certificate_number TEXT,
            died_in_hospital INTEGER
        )
    ''')
    conn.commit()
    conn.close()

def save_to_database(entries, comboboxes, checkboxes):
    conn = sqlite3.connect('forensic_form.db')
    cur = conn.cursor()
    cur.execute('''
        INSERT INTO forensic_data (surname, name, patronymic, gender, birthdate, passport_number, residence, age, nationality,
            death_date, body_received_date, body_received_time, body_found_district, body_comment, examiner, examination_start_date,
            examination_end_date, diagnosis, complications, background_pathology, death_type, death_kind, conclusion_receiver,
            conclusion_issue_date, death_certificate_number, died_in_hospital)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (
        entries['Фамилия'].get(), entries['Имя'].get(), entries['Отчество'].get(), comboboxes['gender'].get(),
        entries['Дата Рождения'].get(), entries['Номер Паспорта'].get(), entries['Место Жительства или Прописки'].get(), entries['Полных лет'].get(),
        entries['Национальность'].get(), entries['Дата Смерти'].get(), entries['Дата доставки трупа'].get(), entries['Время доставки трупа'].get(),
        comboboxes['body_found_district'].get(), entries['Комментарий трупа при нахождении'].get(), comboboxes['examiner'].get(),
        entries['Дата начала экспертизы'].get(), entries['Дата окончания экспертизы'].get(), entries['Диагноз'].get(), entries['Осложнения'].get(),
        entries['Фон патолог'].get(), entries['Вид смерти'].get(), entries['Род смерти'].get(), entries['Кому передано заключение'].get(),
        entries['Дата выдачи заключения'].get(), entries['Номер свидетельства смерти'].get(), checkboxes['died_in_hospital'].get()
    ))
    conn.commit()
    conn.close()
    messagebox.showinfo("Сохранено", "Данные успешно сохранены в базу данных!")

def save_to_word(entries, comboboxes):
    doc = Document()
    for label, entry in entries.items():
        doc.add_paragraph(f"{label}: {entry.get()}")
    
    for label, combobox in comboboxes.items():
        selected_value = combobox.get()
        if selected_value:
            doc.add_paragraph(f"{label}: {selected_value}")
        else:
            messagebox.showwarning("Предупреждение", f"Поле '{label}' не заполнено")
            return
    
    doc.save('output.docx')
    messagebox.showinfo("Успех", "Данные сохранены в файле 'output.docx'")
    print("Документ успешно сохранен")

root = tk.Tk()
root.title("Форма судебно-медицинского эксперта")  

personal_frame = ttk.Frame(root, padding="3 3 12 12")
death_frame = ttk.Frame(root, padding="3 3 12 12")
conclusion_frame = ttk.Frame(root, padding="3 3 12 12")
chemistry_frame = ttk.Frame(root, padding="3 3 12 12")
biology_frame = ttk.Frame(root, padding="3 3 12 12")

ttk.Label(root, text="Персональные данные", font=("Helvetica", 14, "bold")).grid(row=0, column=0)
ttk.Label(root, text="Данные о смерти", font=("Helvetica", 14, "bold")).grid(row=0, column=1)
ttk.Label(root, text="Заключение", font=("Helvetica", 14, "bold")).grid(row=0, column=2)
ttk.Label(root, text="Химия", font=("Helvetica", 14, "bold")).grid(row=3, column=0)  
ttk.Label(root, text="Биология", font=("Helvetica", 14, "bold")).grid(row=3, column=1)  

personal_frame.grid(row=2, column=0, sticky=(tk.N, tk.W, tk.E, tk.S))
death_frame.grid(row=2, column=1, sticky=(tk.N, tk.W, tk.E, tk.S))
conclusion_frame.grid(row=2, column=2, sticky=(tk.N, tk.W, tk.E, tk.S))
chemistry_frame.grid(row=4, column=0, sticky=(tk.N, tk.W, tk.E, tk.S))  
biology_frame.grid(row=4, column=1, sticky=(tk.N, tk.W, tk.E, tk.S)) 

def validate_only_letters(P):
    return P.isalpha() or P == ""

def validate_only_numbers(P):
    return P.isdigit() or P == ""

vcmd_letters = root.register(validate_only_letters)
vcmd_numbers = root.register(validate_only_numbers)

entries = {}
def add_entries(frame, labels, row_start=0):
    for idx, label in enumerate(labels):
        ttk.Label(frame, text=label).grid(row=idx + row_start, column=0, padx=10, pady=5, sticky="w")
        if "Дата" in label or "Номер" in label or "Полных лет" in label or "Время" in label:
            entry = ttk.Entry(frame, validate="key", validatecommand=(vcmd_numbers, '%P'))
            if "Дата" in label:
                entry.bind('<KeyRelease>', format_date)
        else:
            entry = ttk.Entry(frame, validate="key", validatecommand=(vcmd_letters, '%P'))
        entry.grid(row=idx + row_start, column=1, padx=10, pady=5, sticky="ew")
        entries[label] = entry

def format_date(event):
    """Форматирует ввод в формат даты DD/MM/YYYY."""
    content = event.widget.get().replace('/', '')  
    new_content = ''
    if len(content) > 8:
        content = content[:8]  
    if len(content) <= 10:
        new_content = content
    elif 10 < len(content) <= 4:
        new_content = content[:10] + '/' + content[10:]
    elif len(content) > 4:
        new_content = content[:10] + '/' + content[10:4] + '/' + content[4:]

    event.widget.delete(0, tk.END)  
    event.widget.insert(0, new_content) 

labels_personal = ["Фамилия", "Имя", "Отчество", "Пол", "Дата Рождения", "Номер Паспорта",
                   "Место Жительства или Прописки", "Полных лет",
                   "Национальность"]
add_entries(personal_frame, labels_personal)

labels_death = ["Дата Смерти", "Дата доставки трупа", "Время доставки трупа",
                "Район нахождения трупа", "Комментарий трупа при нахождении"]
add_entries(death_frame, labels_death)

labels_conclusion = ["Экспертизу провел", "Дата начала экспертизы", "Дата окончания экспертизы",
                     "Диагноз", "Осложнения", "Фон патолог", "Вид смерти", "Род смерти",
                     "Кому передано заключение", "Дата выдачи заключения", "Номер свидетельства смерти"]
add_entries(conclusion_frame, labels_conclusion)

gender_var = tk.StringVar()
gender_combobox = ttk.Combobox(personal_frame, textvariable=gender_var, values=('М', 'Ж'))
gender_combobox.grid(row=3, column=1, sticky="ew", padx=10, pady=5)

sudmed_var = tk.StringVar()
sudmed_combobox = ttk.Combobox(conclusion_frame, textvariable=sudmed_var, values=('Мағжан Бердібек', 'Оспанов Серік'))
sudmed_combobox.grid(row=0, column=1, sticky="ew", padx=10, pady=5)

rayon_var = tk.StringVar()
rayon_combobox = ttk.Combobox(death_frame, textvariable=rayon_var, values=('Алатауский', 'Алмалинский', 'Ауэзовский', 'Бостандыкский', 'Жетысуский', 'Медеуский', 'Наурызбайский', 'Турксибский'))
rayon_combobox.grid(row=3, column=1, sticky="ew", padx=10, pady=5)

under_one_year_var = tk.BooleanVar()
under_one_year_check = ttk.Checkbutton(personal_frame, text="До 1 года", variable=under_one_year_var)
under_one_year_check.grid(row=len(labels_personal), column=1, sticky="w")

under_fourteen_years_var = tk.BooleanVar()
under_fourteen_years_check = ttk.Checkbutton(personal_frame, text="До 14 лет", variable=under_fourteen_years_var)
under_fourteen_years_check.grid(row=len(labels_personal) + 1, column=1, sticky="w")

died_in_hospital_var = tk.BooleanVar()
died_in_hospital_check = ttk.Checkbutton(death_frame, text="Умер в стационаре", variable=died_in_hospital_var)
died_in_hospital_check.grid(row=len(labels_death), column=1, sticky="w")

ttk.Label(chemistry_frame, text="На Алкоголь", font=("Helvetica", 12, "bold")).grid(row=0, column=0, padx=10, pady=5, sticky="w")
ttk.Label(biology_frame, text="Кровь на группу", font=("Helvetica", 12, "bold")).grid(row=0, column=1, padx=10, pady=5, sticky="w")

alcohol_blood_var = tk.BooleanVar()
alcohol_blood_check = ttk.Checkbutton(chemistry_frame, text="Кровь", variable=alcohol_blood_var)
alcohol_blood_check.grid(row=5, column=0, padx=10, pady=5, sticky="w")

alcohol_urine_var = tk.BooleanVar()
alcohol_urine_check = ttk.Checkbutton(chemistry_frame, text="Моча", variable=alcohol_urine_var)
alcohol_urine_check.grid(row=6, column=0, padx=10, pady=5, sticky="w")

alcohol_muscle_var = tk.BooleanVar()
alcohol_muscle_check = ttk.Checkbutton(chemistry_frame, text="Фрагмент мышц", variable=alcohol_muscle_var)
alcohol_muscle_check.grid(row=7, column=0, padx=10, pady=5, sticky="w")

dry_blood_var = tk.BooleanVar()
dry_blood_check = ttk.Checkbutton(biology_frame, text="Сухая кровь", variable=dry_blood_var)
dry_blood_check.grid(row=1, column=1, padx=10, pady=5, sticky="w")

liquid_blood_var = tk.BooleanVar()
liquid_blood_check = ttk.Checkbutton(biology_frame, text="Жидкая кровь", variable=liquid_blood_var)
liquid_blood_check.grid(row=2, column=1, padx=10, pady=5, sticky="w")

root_teeth_var = tk.BooleanVar()
root_teeth_check = ttk.Checkbutton(biology_frame, text="Коренные зубы", variable=root_teeth_var)
root_teeth_check.grid(row=3, column=1, padx=10, pady=5, sticky="w")

hair_strand_var = tk.BooleanVar()
hair_strand_check = ttk.Checkbutton(biology_frame, text="Прядь волос", variable=hair_strand_var)
hair_strand_check.grid(row=4, column=1, padx=10, pady=5, sticky="w")

def submit_form():
    messagebox.showinfo("Информация", "Форма отправлена")

root.columnconfigure(1, weight=1)

gender_var = tk.StringVar()
gender_combobox = ttk.Combobox(personal_frame, textvariable=gender_var, values=('М', 'Ж'))
gender_combobox.grid(row=3, column=1, sticky="ew", padx=10, pady=5)

sudmed_var = tk.StringVar()
sudmed_combobox = ttk.Combobox(conclusion_frame, textvariable=sudmed_var, values=('Мағжан Бердібек', 'Оспанов Серік'))
sudmed_combobox.grid(row=0, column=1, sticky="ew", padx=10, pady=5)

rayon_var = tk.StringVar()
rayon_combobox = ttk.Combobox(death_frame, textvariable=rayon_var, values=('Алатауский', 'Алмалинский', 'Ауэзовский', 'Бостандыкский', 'Жетысуский', 'Медеуский', 'Наурызбайский', 'Турксибский'))
rayon_combobox.grid(row=3, column=1, sticky="ew", padx=10, pady=5)

comboboxes = {
    "Пол": gender_combobox,
    "Экспертизу Провел": sudmed_combobox,
    "Район нахождения трупа": rayon_combobox
}

save_button = ttk.Button(root, text="Сохранить в Word", command=lambda: save_to_word(entries, comboboxes))
save_button.grid(row=15, column=1, pady=10)

save_button = ttk.Button(root, text="Сохранить в Базу Данных", command=lambda: save_to_database(entries, {'gender': gender_combobox, 'body_found_district': rayon_combobox, 'examiner': sudmed_combobox}, {'died_in_hospital': died_in_hospital_var}))
save_button.grid(row=16, column=1)  

root.mainloop() 
